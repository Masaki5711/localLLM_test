"""Document parser supporting PDF and Word formats."""
import io
import logging
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, str | int | None]
    pages: list[str]


def parse_pdf(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    full_text_parts: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append(text)
            full_text_parts.append(text)

    doc.close()

    return ParsedDocument(
        text="\n\n".join(full_text_parts),
        metadata={
            "file_name": file_name,
            "file_type": "pdf",
            "page_count": len(pages),
            "title": None,
        },
        pages=pages,
    )


def parse_docx(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """Extract text from Word document."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs: list[str] = []
    current_section: list[str] = []
    pages: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_section:
                pages.append("\n".join(current_section))
                current_section = []
            continue

        # Check if heading
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            if current_section:
                pages.append("\n".join(current_section))
                current_section = []
            current_section.append(f"## {text}")
        else:
            current_section.append(text)
        paragraphs.append(text)

    if current_section:
        pages.append("\n".join(current_section))

    # Also extract tables as markdown
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md:
            paragraphs.append(table_md)
            pages.append(table_md)

    return ParsedDocument(
        text="\n\n".join(paragraphs),
        metadata={
            "file_name": file_name,
            "file_type": "docx",
            "page_count": len(pages),
            "title": None,
        },
        pages=pages,
    )


def _table_to_markdown(table) -> str:  # type: ignore[no-untyped-def]
    """Convert a docx table to markdown format."""
    rows: list[str] = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows) if rows else ""


def parse_document(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """Parse a document based on file extension."""
    ext = file_name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes, file_name)
    elif ext in ("docx", "doc"):
        return parse_docx(file_bytes, file_name)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
